"""설계 에이전트 — Spec을 분석하여 Task 목록으로 분해합니다."""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

from claude_agent_sdk import query, ClaudeAgentOptions

from app.models import Spec, Project, Task
from app.repositories import spec_repository, project_repository
from app.websocket import ws_manager
from app.utils.db_handler_sqlalchemy import db_conn
from app.agents.prompts import load_prompt, load_text
from app.agents.guidemap_agent import guidemap_exists, get_design_context

# 가이드라인 파일 경로 (agents/guidemap/PYTHON_FASTAPI_BACKEND_GUIDELINE.md)
_GUIDELINE_PATH = Path(__file__).parent / "guidemap" / "PYTHON_FASTAPI_BACKEND_GUIDELINE.md"

logger = logging.getLogger(__name__)

# Structured Output — 설계 에이전트가 반환해야 하는 JSON 스키마
TASK_LIST_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "analysis_summary": {
            "type": "string",
            "description": "스펙 전체에 대한 분석 요약",
        },
        "tasks": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "description": {"type": "string"},
                    "acceptance_criteria": {
                        "type": "array",
                        "items": {"type": "string"},
                    },
                    "target_files": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "File paths to create or modify, relative to the project root (e.g. app/api/users.py)",
                    },
                    "priority": {
                        "type": "string",
                        "enum": ["low", "medium", "high", "critical"],
                    },
                    "complexity": {
                        "type": "string",
                        "enum": ["trivial", "low", "medium", "high", "very_high"],
                    },
                },
                "required": [
                    "title",
                    "description",
                    "acceptance_criteria",
                    "priority",
                    "complexity",
                ],
            },
        },
        "db_schema": {
            "type": "object",
            "description": "DB schema design for new projects only. Omit for existing projects.",
            "properties": {
                "overview": {
                    "type": "string",
                    "description": "Overall description of the DB schema",
                },
                "tables": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "description": {"type": "string"},
                            "columns": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "name": {"type": "string"},
                                        "type": {"type": "string"},
                                        "nullable": {"type": "boolean"},
                                        "description": {"type": "string"},
                                    },
                                    "required": ["name", "type", "nullable"],
                                },
                            },
                        },
                        "required": ["name", "columns"],
                    },
                },
            },
            "required": ["overview", "tables"],
        },
    },
    "required": ["analysis_summary", "tasks"],
}


def _get_stack_context(project: Project) -> str:

    # 파이썬 + fastapi의 경우    
    if project.project_stack == "python" and project.framework == "fastapi":
        # 신규 프로젝트의 경우 fastapi_new_project.md의 규율을 따른다()
        if project.project_type == "new":
            summary = load_text("fastapi_new_project.md")
            return (
                f"- This is a new FastAPI project. Review the key rules summary below,"
                f"{summary}"
            )
            
        # 기존 프로젝트의 경우, 별도로 생성된 guideline을 참고합니다
        if guidemap_exists(project.name):
            design_context = get_design_context(project.name)
            return (
                f"- This is an existing FastAPI project.\n"
                f"- Follow the project guide below:\n\n"
                f"{design_context}\n"
            )
        return (
            "- This is an existing FastAPI project\n"
            "- Maintain the existing layer structure (router → service → repository)\n"
            "- Follow the existing response format and exception handling patterns\n"
        )
    if project.project_stack == "java":
        return (
            "- This project uses an existing internal framework\n"
            "- Follow the existing package structure and naming conventions\n"
        )
    return ""


def _build_prompt(spec_content: str, project: Project) -> str:
    stack_context = _get_stack_context(project)

    codebase_section = ""
    if project.local_repo_path:
        guidemap_ready = (
            project.project_type == "existing"
            and guidemap_exists(project.name)
        )
        if not guidemap_ready:
            codebase_section = load_prompt(
                "codebase_section.md",
                local_repo_path=project.local_repo_path,
            )

    return load_prompt(
        "design_agent.md",
        project_name=project.name,
        project_stack=project.project_stack,
        framework=project.framework or "unspecified",
        repo_url=project.repo_url or "not configured",
        stack_context=stack_context,
        codebase_section=codebase_section,
        spec_content=spec_content,
    )


def _extract_docx_text(path: str) -> str:
    """Extracts plain text from a .docx file using python-docx."""
    import docx  # python-docx

    doc = docx.Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


async def _load_spec_content(spec: Spec) -> str:
    """Extracts text content from a Spec for analysis."""
    if spec.raw_content:
        return spec.raw_content

    if spec.source_path:
        path_lower = spec.source_path.lower()
        try:
            if path_lower.endswith(".docx") or path_lower.endswith(".doc"):
                return _extract_docx_text(spec.source_path)
            # Plain text files
            with open(spec.source_path, encoding="utf-8", errors="ignore") as f:
                return f.read()
        except OSError as e:
            logger.warning("Failed to read spec file: %s", e)
            return f"[File read error: {spec.source_path}]"

    return "[No spec content]"


async def _update_spec_status(spec_id: str, status: str, analysis_result: str | None = None) -> None:
    """Updates the spec status and analysis result."""
    async with db_conn.transaction() as session:
        spec = await spec_repository.find_by_id(spec_id, session)
        if spec:
            spec.status = status
            if analysis_result is not None:
                spec.analysis_result = analysis_result
            await session.flush()


def _save_db_schema_file(local_repo_path: str, db_schema: dict) -> Path:
    """Writes the DB schema to {local_repo_path}/docs/schema/db_schema.md."""
    schema_dir = Path(local_repo_path) / "docs" / "schema"
    schema_dir.mkdir(parents=True, exist_ok=True)
    schema_file = schema_dir / "db_schema.md"

    lines = ["# DB Schema\n", f"\n{db_schema.get('overview', '')}\n"]
    for table in db_schema.get("tables", []):
        lines.append(f"\n## {table['name']}\n")
        if table.get("description"):
            lines.append(f"{table['description']}\n")
        lines.append("\n| Column | Type | Nullable | Description |\n")
        lines.append("|--------|------|----------|-------------|\n")
        for col in table.get("columns", []):
            nullable = "O" if col.get("nullable") else "X"
            lines.append(
                f"| {col['name']} | {col['type']} | {nullable} | {col.get('description', '')} |\n"
            )

    schema_file.write_text("".join(lines), encoding="utf-8")
    logger.info("DB schema written to %s", schema_file)
    return schema_file


async def _save_tasks(project_id: str, spec_id: str, task_items: list[dict]) -> list[Task]:
    """Saves the decomposed Task list to the database."""
    tasks = [
        Task(
            project_id=project_id,
            spec_id=spec_id,
            title=item["title"],
            description=item["description"],
            acceptance_criteria=item.get("acceptance_criteria"),
            target_files=item.get("target_files"),
            priority=item.get("priority", "medium"),
            complexity=item.get("complexity", "medium"),
            status="plan_reviewing",
            sort_order=idx,
        )
        for idx, item in enumerate(task_items)
    ]
    async with db_conn.transaction() as session:
        for task in tasks:
            session.add(task)
        await session.flush()
        for task in tasks:
            await session.refresh(task)
    return tasks


async def analyze_spec_and_create_tasks(spec_id: str) -> None:
    """
    Analyzes a Spec and auto-generates a Task list.

    1. Load Spec → set status to 'analyzing'
    2. Run Design Agent (claude-agent-sdk query)
    3. Save analysis result to DB, set Spec status to 'analyzed'
    4. Broadcast progress via WebSocket at each step
    """
    # 1. Load Spec / Project
    spec = await spec_repository.find_by_id(spec_id)
    if not spec:
        logger.error("Spec not found: %s", spec_id)
        return

    project = await project_repository.find_by_id(spec.project_id)
    if not project:
        logger.error("Project not found: %s", spec.project_id)
        return

    project_id = project.id

    async def broadcast(msg: dict) -> None:
        await ws_manager.broadcast(project_id, msg)

    # 2. Set status → analyzing
    await _update_spec_status(spec_id, "analyzing")
    await broadcast({"type": "spec_analyzing", "spec_id": spec_id})

    try:
        spec_content = await _load_spec_content(spec)
        prompt = _build_prompt(spec_content, project)

        options = ClaudeAgentOptions(
            model="claude-opus-4-6",
            allowed_tools=["Read", "Glob", "Grep"],
            permission_mode="bypassPermissions",
            max_turns=10,
            output_format={"type": "json_schema", "schema": TASK_LIST_SCHEMA},
        )

        parsed: dict | None = None

        async for message in query(prompt=prompt, options=options):
            # Forward agent messages to WebSocket
            try:
                msg_data = message.model_dump() if hasattr(message, "model_dump") else str(message)
            except Exception:
                msg_data = str(message)

            await broadcast({"type": "agent_message", "spec_id": spec_id, "message": msg_data})

            # Extract structured_output from ResultMessage (when output_format=json_schema)
            if hasattr(message, "structured_output") and message.structured_output is not None:
                parsed = message.structured_output

            # Fallback: parse JSON from AssistantMessage result text
            elif hasattr(message, "result") and message.result:
                try:
                    parsed = json.loads(message.result)
                except (json.JSONDecodeError, TypeError):
                    pass

        # 3. Parse result
        if not parsed:
            raise ValueError("Agent returned no result.")
        task_items: list[dict] = (parsed or {}).get("tasks", [])
        analysis_summary: str = (parsed or {}).get("analysis_summary", "")

        # 4a. Save DB schema file for new projects
        db_schema = (parsed or {}).get("db_schema")
        if project.project_type == "new" and db_schema and project.local_repo_path:
            _save_db_schema_file(project.local_repo_path, db_schema)

        # 4b. Save Tasks to DB
        saved_tasks = await _save_tasks(project_id, spec_id, task_items)

        # 5. Set Spec status → analyzed
        await _update_spec_status(spec_id, "analyzed")

        # 6. Broadcast completion
        await broadcast(
            {
                "type": "spec_analyzed",
                "spec_id": spec_id,
                "analysis_summary": analysis_summary,
                "tasks": [
                    {
                        "id": t.id,
                        "project_id": t.project_id,
                        "spec_id": t.spec_id,
                        "title": t.title,
                        "description": t.description,
                        "priority": t.priority,
                        "complexity": t.complexity,
                        "status": t.status,
                    }
                    for t in saved_tasks
                ],
            }
        )

    except Exception as exc:
        logger.exception("Design Agent failed (spec_id=%s): %s", spec_id, exc)
        await _update_spec_status(spec_id, "uploaded")  # Revert status on failure
        await broadcast(
            {
                "type": "spec_analyze_failed",
                "spec_id": spec_id,
                "error": str(exc),
            }
        )
