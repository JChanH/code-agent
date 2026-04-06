"""설계 에이전트 v2 — 2단계 방식.

1단계: Claude API + agentic loop로 자유 분석 (read/glob/grep 도구 사용)
2단계: anthropic API + tool_choice로 JSON 포맷팅
"""

from __future__ import annotations

import logging
from typing import Any

import anthropic

from app.models import Spec, Project, Task
from app.repositories import spec_repository, project_repository
from app.websocket import make_broadcaster
from app.websocket.messages import (
    msg_agent_message,
    msg_spec_analyzed,
    msg_spec_analyze_failed,
    msg_spec_analyzing,
)
from app.utils.db_handler_sqlalchemy import db_conn
from app.agents.prompts import load_prompt
from app.agents.guidemap_agent import guidemap_exists, get_guidemap_context
from app.agents.tools.agent_loop import run_agent_loop
from app.config import get_settings

logger = logging.getLogger(__name__)

# Structured Output 스키마 — 포맷팅 단계에서 tool로 사용
TASK_LIST_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "tasks": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "description": {"type": "string"},
                    "acceptance_criteria": {
                        "type": "array",
                        "items": {"type": "string"},
                    },
                    "implementation_steps": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Step-by-step implementation plan for the code agent.",
                    },
                },
                "required": [
                    "title",
                    "description",
                    "acceptance_criteria",
                    "implementation_steps",
                ],
            },
        },
    },
    "required": ["tasks"],
}

# 포맷팅 단계에서 강제할 tool 정의
_FORMAT_TOOL: dict[str, Any] = {
    "name": "create_tasks",
    "description": "분석 결과를 Task 목록으로 구조화합니다.",
    "input_schema": TASK_LIST_SCHEMA,
}


def _get_stack_context(project: Project, use_guidemap: bool = True) -> str:
    if use_guidemap and guidemap_exists(project.name):
        design_context = get_guidemap_context(project.name)
        return (
            f"- This is an existing project.\n"
            f"- Follow the project guide below:\n\n"
            f"{design_context}\n"
        )
    return ""


def _build_prompt(spec_content: str, project: Project, use_guidemap: bool = True) -> str:
    stack_context = _get_stack_context(project, use_guidemap=use_guidemap)

    codebase_section = ""
    if project.local_repo_path:
        guidemap_ready = (
            use_guidemap and
            project.project_type == "existing" and
            guidemap_exists(project.name)
        )
        
        # guidemap이 없는 경우에만 전체 탐색
        if not guidemap_ready:
            codebase_section = load_prompt(
                "codebase_section.md",
                local_repo_path=project.local_repo_path,
            )

    return load_prompt(
        "design_agent.md",
        project_name=project.name,
        project_stack=project.project_stack,
        framework=project.framework or "unspecified",
        repo_url=project.repo_url or "not configured",
        stack_context=stack_context,
        codebase_section=codebase_section,
        spec_content=spec_content,
    )


def _extract_docx_text(path: str) -> str:
    import docx

    doc = docx.Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


async def _load_spec_content(spec: Spec) -> str:
    if spec.raw_content:
        return spec.raw_content

    if spec.source_path:
        path_lower = spec.source_path.lower()
        try:
            if path_lower.endswith(".docx") or path_lower.endswith(".doc"):
                return _extract_docx_text(spec.source_path)
            with open(spec.source_path, encoding="utf-8", errors="ignore") as f:
                return f.read()
        except OSError as e:
            logger.warning("Failed to read spec file: %s", e)
            return f"[File read error: {spec.source_path}]"

    return "[No spec content]"


async def _update_spec_status(spec_id: str, status: str, analysis_result: str | None = None) -> None:
    async with db_conn.transaction() as session:
        spec = await spec_repository.find_by_id(spec_id, session)
        if spec:
            spec.status = status
            if analysis_result is not None:
                spec.analysis_result = analysis_result
            await session.flush()


async def _save_tasks(project_id: str, spec_id: str, task_items: list[dict]) -> list[Task]:
    tasks = [
        Task(
            project_id=project_id,
            spec_id=spec_id,
            title=item["title"],
            description=item["description"],
            acceptance_criteria=item.get("acceptance_criteria"),
            implementation_steps=item.get("implementation_steps"),
            status="plan_reviewing",
        )
        for item in task_items
    ]
    async with db_conn.transaction() as session:
        for task in tasks:
            session.add(task)
        await session.flush()
        for task in tasks:
            await session.refresh(task)
    return tasks


def _format_with_anthropic(free_text: str) -> dict:
    """에이전트 자유 분석 결과를 Anthropic API + tool_choice로 JSON 포맷팅."""
    settings = get_settings()
    client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4096,
        tools=[_FORMAT_TOOL],
        tool_choice={"type": "tool", "name": "create_tasks"},
        messages=[
            {
                "role": "user",
                "content": (
                    "아래 분석 결과를 create_tasks tool을 사용해 Task 목록으로 구조화하세요.\n\n"
                    f"<analysis>\n{free_text}\n</analysis>"
                ),
            }
        ],
    )

    for block in response.content:
        if block.type == "tool_use" and block.name == "create_tasks":
            return block.input

    raise ValueError("Formatting step returned no tool_use block.")


async def analyze_spec_and_create_tasks(spec_id: str, use_guidemap: bool = False) -> None:
    """
    Spec을 분석하고 Task 목록을 자동 생성합니다.

    1단계: claude-agent-sdk로 자유 분석 (output_format 없음)
    2단계: anthropic API + tool_choice로 JSON 포맷팅

    Args:
        spec_id: 분석할 Spec ID
        use_guidemap: False로 설정하면 guidemap을 무시하고 분석합니다.
    """
    # spec + project 조회
    spec = await spec_repository.find_by_id(spec_id)
    if not spec:
        logger.error("Spec not found: %s", spec_id)
        return

    project = await project_repository.find_by_id(spec.project_id)
    if not project:
        logger.error("Project not found: %s", spec.project_id)
        return

    project_id = project.id
    broadcast = make_broadcaster(project_id)

    # 스펙서 업데이트
    await _update_spec_status(spec_id, "analyzing")
    await broadcast(msg_spec_analyzing(spec_id))

    try:
        spec_content = await _load_spec_content(spec)
        prompt = _build_prompt(
            spec_content,
            project,
            use_guidemap=use_guidemap
        )

        # 1단계: agentic loop로 자유 분석
        settings = get_settings()
        client = anthropic.AsyncAnthropic(api_key=settings.anthropic_api_key)

        collected_texts: list[str] = []

        async def on_message(payload: dict) -> None:
            await broadcast(msg_agent_message(payload, spec_id=spec_id))
            # 텍스트 블록 수집 (최종 분석 결과 취합용)
            for item in payload.get("content", []):
                if item.get("type") == "text":
                    collected_texts.append(item["text"])

        final_text, _ = await run_agent_loop(
            client=client,
            model="claude-sonnet-4-6",
            messages=[{"role": "user", "content": prompt}],
            tool_names=["read_file", "glob_files", "grep_search"],
            max_turns=30,
            working_dir=project.local_repo_path,
            on_message=on_message,
        )

        free_text = final_text or "\n".join(collected_texts)
        if not free_text.strip():
            raise ValueError("Agent returned no analysis result.")

        # 2단계: Anthropic API로 JSON 포맷팅
        parsed = _format_with_anthropic(free_text)
        task_items: list[dict] = parsed.get("tasks", [])

        saved_tasks = await _save_tasks(project_id, spec_id, task_items)

        await _update_spec_status(spec_id, "analyzed")

        await broadcast(
            msg_spec_analyzed(
                spec_id=spec_id,
                analysis_summary="",
                tasks=[
                    {
                        "id": t.id,
                        "project_id": t.project_id,
                        "spec_id": t.spec_id,
                        "title": t.title,
                        "description": t.description,
                        "status": t.status,
                    }
                    for t in saved_tasks
                ],
            )
        )

    except Exception as exc:
        logger.exception("Design Agent v2 failed (spec_id=%s): %s", spec_id, exc)
        await _update_spec_status(spec_id, "uploaded")
        await broadcast(msg_spec_analyze_failed(spec_id, str(exc)))